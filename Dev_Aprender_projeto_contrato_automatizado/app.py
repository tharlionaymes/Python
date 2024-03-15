#1 - Anotar o passo a passo de como pretende resolver essa demanda

#2 - Quais técnicas podem resolver esse problema
#2.1 - Ler planilhas: openpyxl
#2.2 - Criar arquivos word: python-docx

from openpyxl import load_workbook #"load_workbook" permite carregar a pasta de trabalho
from docx import Document #"Document" permite criar arquivos word
from datetime import datetime # "datetime" importa a data atual

#3 - Passar as informações da planilha para o arquivo word

# Gravar em uma variável o carregamento da pasta de trabalho
planilha_fornecedores = load_workbook('/home/vermilion/Python/Dev_Aprender_projeto_contrato_automatizado/fornecedores.xlsx') 

# Seleciona a planilha que queremos extrair as informações
pagina_fornecedores = planilha_fornecedores["Sheet1"]

# Percorrer cada linha da planilha e extrair as informações.
# "iter_rows" - Percorrer linhas. "(min_row=2)" - Qual linha que começara a percorrer, no caso a segunda linha - "values_only=true" - Seleciona apenas os dados(texto)
for linha in pagina_fornecedores.iter_rows(min_row=2, values_only=True):
    nome_empresa, endereco, cidade, estado, cep, telefone, email, setor  = linha
    
    # Criar o arquivo word
    arquivo_word = Document()

    # Criar um cabeçalho (vamos colocar aquele que esta no contrato mesmo)
    arquivo_word.add_heading('Contrato de Prestação de Serviço', 0)

    texto_contrato = f"""

    Este contrato de prestação de serviços é feito entre {nome_empresa}, com endereço em {endereco},
    {cidade}, {estado}, CEP {cep}, doravante denominado FORNECEDOR, e a empresa CONTRATANTE.

    Pelo presente instrumento particular, as partes têm, entre si, justo e acordado o seguinte:

    1. OBJETO DO CONTRATO
    O FORNECEDOR compromete-se a fornecer à CONTRATANTE os serviços/material de acordo com as especificações acordadas, respeitando os padrões de qualidade e os prazos estipulados.

    2. PRAZO
    Este contrato tem prazo de vigência de 12 (doze) meses, iniciando-se na data de sua assinatura, podendo ser renovado conforme acordo entre as partes.

    3. VALOR E FORMA DE PAGAMENTO
    O valor dos serviços prestados será acordado conforme as demandas da CONTRATANTE e a capacidade de entrega do FORNECEDOR. Os pagamentos serão realizados mensalmente, mediante apresentação de nota fiscal.

    4. CONFIDENCIALIDADE
    Todas as informações trocadas entre as partes durante a vigência deste contrato serão tratadas como confidenciais.

    Para firmeza e como prova de assim haverem justo e contratado, as partes assinam o presente contrato em duas vias de igual teor e forma.

    FORNECEDOR: {nome_empresa}
    E-mail: {email}

    CONTRATANTE: Prestadores S/A 
    E-mail: prestadores_sa@gmail.com

    Rio de Janeiro,{datetime.now().strftime('%d/%m/%Y')}

    """
    # Inserir o texto do contrato ao arquivo word
    arquivo_word.add_paragraph(texto_contrato)

    #4 - Salvar aquele arquivo word em uma pasta específica chamada "contratos"
    arquivo_word.save(f"/home/vermilion/Python/Dev_Aprender_projeto_contrato_automatizado/contratos/contrato_{nome_empresa}.docx")


#5 - Repetir para todas as linhas da planilha, ou seja, automatizar
# Isso é feito com o for


